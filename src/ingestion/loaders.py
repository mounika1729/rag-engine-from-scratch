"""
src/ingestion/loaders.py — Load raw content from various file types.

Supported:
  • Plain text  (.txt)
  • PDF         (.pdf)   — via pypdf
  • Word doc    (.docx)  — via python-docx
  • Markdown    (.md)
  • Web URL     (http/https) — via requests + BeautifulSoup
  • Directory   — recursively loads all supported files
"""

from __future__ import annotations

import re
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from loguru import logger

from src.models import Document


# ── helpers ───────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Collapse whitespace and strip leading/trailing blanks."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── individual loaders ────────────────────────────────────────────────────────

def load_text(path: Path) -> Document:
    content = path.read_text(encoding="utf-8", errors="ignore")
    return Document(
        content=_clean(content),
        metadata={"source": str(path), "file_type": "txt"},
    )


def load_pdf(path: Path) -> Document:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("Install pypdf: pip install pypdf")

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    content = "\n\n".join(pages)
    return Document(
        content=_clean(content),
        metadata={
            "source": str(path),
            "file_type": "pdf",
            "num_pages": len(reader.pages),
        },
    )


def load_docx(path: Path) -> Document:
    try:
        from docx import Document as DocxDoc
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")

    doc = DocxDoc(str(path))
    content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return Document(
        content=_clean(content),
        metadata={"source": str(path), "file_type": "docx"},
    )


def load_url(url: str, timeout: int = 15) -> Document:
    headers = {"User-Agent": "RAGBot/1.0"}
    resp = requests.get(url, headers=headers, timeout=timeout)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    # Remove scripts, styles, nav
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    content = soup.get_text(separator="\n")
    return Document(
        content=_clean(content),
        metadata={"source": url, "file_type": "url"},
    )


# ── dispatcher ────────────────────────────────────────────────────────────────

_LOADERS = {
    ".txt": load_text,
    ".md": load_text,
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_document(source: str | Path) -> Document:
    """
    Auto-detect source type and return a Document.

    Args:
        source: File path or HTTP/HTTPS URL string.
    """
    src = str(source)
    if src.startswith("http://") or src.startswith("https://"):
        logger.info(f"Loading URL: {src}")
        return load_url(src)

    path = Path(source)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    ext = path.suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: '{ext}'. Supported: {list(_LOADERS)}")

    logger.info(f"Loading {ext.upper()} file: {path}")
    return loader(path)


def load_directory(directory: str | Path, recursive: bool = True) -> list[Document]:
    """
    Load all supported documents from a directory.

    Args:
        directory: Path to directory.
        recursive: If True, scan subdirectories too.
    """
    dir_path = Path(directory)
    if not dir_path.is_dir():
        raise NotADirectoryError(f"Not a directory: {dir_path}")

    pattern = "**/*" if recursive else "*"
    docs: list[Document] = []

    for path in sorted(dir_path.glob(pattern)):
        if path.suffix.lower() in _LOADERS:
            try:
                doc = load_document(path)
                docs.append(doc)
                logger.debug(f"Loaded: {path.name} ({len(doc.content)} chars)")
            except Exception as exc:
                logger.warning(f"Skipping {path.name}: {exc}")

    logger.info(f"Loaded {len(docs)} documents from '{dir_path}'")
    return docs
